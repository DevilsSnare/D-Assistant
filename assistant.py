import os
import speech_recognition as sr
from datetime import date as dt
import nltk
from nameparser.parser import HumanName
import re
import spacy
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, RGBColor, Inches
from docx.shared import Pt


def create(name, age, date, tablet):
    # init document
    document = Document()
    para = document.add_heading("")
    # Header
    runn = para.add_run("HealthCare Hospitals \n")
    run = para.add_run(
        " Dr.Rishi Raj M.D.(Neurology) \t \t Dr.Swathy M.S.(Ortho) \t Dr.Shreevarshann M.S.(Ophthalmology) \n No.20,Blossom Avenue,Madura \n Ph:0452-2588 1522 \n")
    # Set page attributes
    font = run.font
    runn.font.size = Pt(18)
    runn.font.color.rgb = RGBColor(153, 17, 150)
    font.size = Pt(16)
    font.color.rgb = RGBColor(217, 17, 213)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para1 = document.add_paragraph()
    run1 = para1.add_run("Name:{}".format(name))
    run1.font.size = Pt(14)

    para2 = document.add_paragraph()
    run2 = para2.add_run("Age:{}".format(age))
    run2.font.size = Pt(14)

    para3 = document.add_paragraph()
    run3 = para3.add_run("Date:{}".format(date))
    run3.font.size = Pt(14)
    table = document.add_table(len(tablet), 5)

    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Name'
    heading_cells[1].text = 'Dosage and Strength'
    heading_cells[2].text = 'Duration and Frequency'

    for i in range(len(tablet)):
        cells = table.add_row().cells
        cells[0].text = tablet[i][0]
        cells[1].text = tablet[i][1]
        cells[2].text = tablet[i][2]

    para5 = document.add_paragraph("\n \n")

    header = document.sections[0].footer
    f1 = header.add_paragraph()

    runf = f1.add_run(
        "\n Signature \n I hereby accept that this prescription was verified")
    runf.font.size = Pt(16)
    runf.font.color.rgb = RGBColor(0, 0, 0)
    f1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.save("{}.docx".format(name))


def extract_details(text):
    med7 = spacy.load("en_core_med7_lg")

    # create distinct colours for labels
    col_dict = {}
    seven_colours = ['#e6194B', '#3cb44b', '#ffe119',
                     '#ffd8b1', '#f58231', '#f032e6', '#42d4f4']
    for label, colour in zip(med7.pipe_labels['ner'], seven_colours):
        col_dict[label] = colour
    options = {'ents': med7.pipe_labels['ner'], 'colors': col_dict}

    doc = med7(text)

    #spacy.displacy.render(doc, style='ent', jupyter=True, options=options)

    return [(ent.text, ent.label_) for ent in doc.ents]


def get_human_names(text):
    # uncomment below line for the first time
    # nltk.download()

    tokens = nltk.tokenize.word_tokenize(text)
    pos = nltk.pos_tag(tokens)
    sentt = nltk.ne_chunk(pos, binary=False)
    person_list = []
    person = []
    name = ""
    for subtree in sentt.subtrees(filter=lambda t: t.label() == 'PERSON'):
        for leaf in subtree.leaves():
            person.append(leaf[0])
        for part in person:
            name += part + ' '
        if name[:-1] not in person_list:
            person_list.append(name[:-1])
        name = ''
        person = []

    return (person_list)


def get_age(text):
    return re.findall(r'\d{1,3}', text)[0]


def extract(text):
    name = " ".join(get_human_names(text))
    age = get_age(text)
    date = str(dt.today())
    details = extract_details(text)
    medicine = []
    i = 0
    while i < len(details):
        while i < len(details) and details[i][1] == "DRUG":
            temp = ["", "", ""]
            temp[0] = details[i][0]
            i += 1
            while i < len(details) and details[i][1] != "DRUG":
                if details[i][1] == "STRENGTH" or details[i][1] == "DOSAGE":
                    temp[1] += details[i][0] + " "

                if details[i][1] == "DURATION" or details[i][1] == "FREQUENCY":
                    temp[2] += details[i][0] + " "

                i += 1
            medicine.append(temp)
        i += 1

    return name, age, date, medicine


def liveRecord():
    allEars = sr.Recognizer()
    with sr.Microphone() as source:
        print("Start Speaking! D-Assistant is all ears.")
        audio = allEars.listen(source)
        print("Listening completed!")

    try:
        text = allEars.recognize_google(audio)
        name, age, date, medicine = extract(text)
        return (text.upper(), name, age, date, medicine)
    except:
        return ("Error while converting speech to text! Refresh page to try again.", 'unknown', 'unknown', str(dt.today()), 'unknown')
